from __future__ import annotations

import io
from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from pathlib import Path
from scripts.synthetic_data.core.document_rendering.context import (
    GeneratedDocument,
    RenderContext,
)
from scripts.synthetic_data.core.font_registry import font_for_script
from scripts.synthetic_data.core.text_image_renderer import render_text_snippet
from scripts.synthetic_data.core.text_shaping import shape_for_display


_EAST_ASIAN_SCRIPTS = frozenset({"Han", "Han/Kana", "Hangul"})


def _mark_paragraph_rtl(paragraph: Paragraph) -> None:
    paragraph_properties = paragraph._p.get_or_add_pPr()
    paragraph_properties.append(
        paragraph_properties.makeelement(qn("w:bidi"), {})
    )


def _apply_run_style(
    run: Run, *, font_name: str, script: str, is_rtl: bool
) -> None:
    run.font.name = font_name
    run_properties = run._r.get_or_add_rPr()
    east_fonts = run_properties.makeelement(qn("w:rFonts"), {})
    east_fonts.set(qn("w:ascii"), font_name)
    east_fonts.set(qn("w:hAnsi"), font_name)
    east_fonts.set(qn("w:cs"), font_name)
    if script in _EAST_ASIAN_SCRIPTS:
        east_fonts.set(qn("w:eastAsia"), font_name)
    run_properties.append(east_fonts)
    if is_rtl:
        run_properties.append(run_properties.makeelement(qn("w:rtl"), {}))


def _add_styled_paragraph(
    document: DocxDocument,
    text: str,
    *,
    font_name: str,
    script: str,
    is_rtl: bool,
) -> Paragraph:
    paragraph = document.add_paragraph()
    if is_rtl:
        _mark_paragraph_rtl(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run(text)
    _apply_run_style(run, font_name=font_name, script=script, is_rtl=is_rtl)
    return paragraph


def _build_copyable_document(ctx: RenderContext) -> DocxDocument:
    font = font_for_script(ctx.script)
    latin_font = font_for_script("Latin")
    document = Document()
    marker_en, marker_local = ctx.marker_lines
    _add_styled_paragraph(
        document,
        marker_en,
        font_name=latin_font.font_name,
        script="Latin",
        is_rtl=False,
    )
    _add_styled_paragraph(
        document,
        marker_local,
        font_name=font.font_name,
        script=ctx.script,
        is_rtl=ctx.is_rtl,
    )
    document.add_heading(ctx.country.name, level=1)
    for block in ctx.recipe.blocks:
        _add_styled_paragraph(
            document,
            block.text,
            font_name=font.font_name,
            script=ctx.script,
            is_rtl=ctx.is_rtl,
        )
    return document


def render_docx_copyable(
    ctx: RenderContext, *, output_dir: Path
) -> GeneratedDocument:
    output_dir.mkdir(parents=True, exist_ok=True)
    path = output_dir / f"{ctx.recipe.recipe_id}.docx"
    _build_copyable_document(ctx).save(str(path))
    return GeneratedDocument(
        path=path,
        file_format="docx",
        mode="copyable",
        locale=ctx.recipe.locale,
        country_id=ctx.country.country_id,
        recipe_id=ctx.recipe.recipe_id,
        related_artifact_ids=(ctx.country.country_id,),
        size_bytes=path.stat().st_size,
    )


def render_docx_non_copyable(
    ctx: RenderContext, *, output_dir: Path
) -> GeneratedDocument:
    output_dir.mkdir(parents=True, exist_ok=True)
    path = output_dir / f"{ctx.recipe.recipe_id}.docx"
    font = font_for_script(ctx.script)
    document = Document()
    for line in (
        *ctx.marker_lines,
        ctx.country.name,
        *(block.text for block in ctx.recipe.blocks),
    ):
        display_text = shape_for_display(line, is_rtl=ctx.is_rtl)
        snippet_image = render_text_snippet(
            display_text, font_path=str(font.file_path)
        )
        image_buffer = io.BytesIO()
        snippet_image.save(image_buffer, format="PNG")
        image_buffer.seek(0)
        document.add_picture(image_buffer)
    document.save(str(path))
    return GeneratedDocument(
        path=path,
        file_format="docx",
        mode="non_copyable",
        locale=ctx.recipe.locale,
        country_id=ctx.country.country_id,
        recipe_id=ctx.recipe.recipe_id,
        related_artifact_ids=(ctx.country.country_id,),
        size_bytes=path.stat().st_size,
    )


def render_docx_mixed(
    ctx: RenderContext, *, output_dir: Path
) -> GeneratedDocument:
    output_dir.mkdir(parents=True, exist_ok=True)
    path = output_dir / f"{ctx.recipe.recipe_id}.docx"
    font = font_for_script(ctx.script)
    document = _build_copyable_document(ctx)
    document.add_heading("Scanned fragment", level=2)
    scanned_text = shape_for_display(
        ctx.recipe.blocks[-1].text, is_rtl=ctx.is_rtl
    )
    snippet_image = render_text_snippet(
        scanned_text, font_path=str(font.file_path)
    )
    image_buffer = io.BytesIO()
    snippet_image.save(image_buffer, format="PNG")
    image_buffer.seek(0)
    document.add_picture(image_buffer)
    document.save(str(path))
    return GeneratedDocument(
        path=path,
        file_format="docx",
        mode="mixed",
        locale=ctx.recipe.locale,
        country_id=ctx.country.country_id,
        recipe_id=ctx.recipe.recipe_id,
        related_artifact_ids=(ctx.country.country_id,),
        size_bytes=path.stat().st_size,
    )
