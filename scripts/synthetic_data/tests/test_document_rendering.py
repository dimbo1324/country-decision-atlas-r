from __future__ import annotations

import json
import pytest
from docx import Document as DocxReader
from openpyxl import load_workbook
from pathlib import Path
from pypdf import PdfReader
from scripts.synthetic_data.core.document_rendering.context import RenderContext
from scripts.synthetic_data.core.document_rendering.docx_renderer import (
    render_docx_copyable,
    render_docx_mixed,
    render_docx_non_copyable,
)
from scripts.synthetic_data.core.document_rendering.json_renderer import (
    render_json,
)
from scripts.synthetic_data.core.document_rendering.pdf_renderer import (
    render_pdf_copyable,
    render_pdf_mixed,
    render_pdf_non_copyable,
)
from scripts.synthetic_data.core.document_rendering.txt_renderer import (
    render_txt,
)
from scripts.synthetic_data.core.document_rendering.xlsx_renderer import (
    render_xlsx,
)
from scripts.synthetic_data.core.locale_corpus import (
    LocaleCorpus,
    load_locale_corpus,
)
from scripts.synthetic_data.core.world_generator import (
    WorldGenerationOptions,
    WorldGenerator,
)
from scripts.synthetic_data.core.world_input import load_world_input
from scripts.synthetic_data.core.world_models import (
    FICTIONAL_NOTICE,
    SyntheticWorld,
)


# A representative subset covering every distinct code path: plain Latin,
# Cyrillic, an RTL script, and a no-space-wrapping (Han) script. The full
# 15-locale sweep is exercised separately (PDF-only) below, keeping this
# suite's overall runtime reasonable.
_SAMPLE_LOCALES = ("en-US", "ru-RU", "ar-SA", "zh-Hans-CN")


def _world() -> SyntheticWorld:
    return WorldGenerator(input_data=load_world_input()).generate(
        WorldGenerationOptions(seed=42017, profile="balanced")
    )


def _corpus() -> LocaleCorpus:
    return load_locale_corpus()


def _context_for(
    world: SyntheticWorld, corpus: LocaleCorpus, locale: str
) -> RenderContext:
    recipe = next(r for r in world.document_recipes if r.locale == locale)
    country = next(
        c for c in world.countries if c.country_id == recipe.country_id
    )
    return RenderContext(
        world=world,
        recipe=recipe,
        country=country,
        text_pack=corpus.by_locale(locale),
    )


@pytest.fixture(scope="module")
def world() -> SyntheticWorld:
    return _world()


@pytest.fixture(scope="module")
def corpus() -> LocaleCorpus:
    return _corpus()


@pytest.mark.parametrize("locale", _SAMPLE_LOCALES)
def test_json_renderer_produces_valid_utf8_json(
    world: SyntheticWorld, corpus: LocaleCorpus, tmp_path: Path, locale: str
) -> None:
    ctx = _context_for(world, corpus, locale)

    document = render_json(ctx, output_dir=tmp_path)
    payload = json.loads(document.path.read_text(encoding="utf-8"))

    assert payload["locale"] == locale
    assert payload["blocks"]
    assert payload["synthetic_marking"]["en"] == FICTIONAL_NOTICE


@pytest.mark.parametrize("locale", _SAMPLE_LOCALES)
def test_txt_renderer_contains_marking_and_blocks(
    world: SyntheticWorld, corpus: LocaleCorpus, tmp_path: Path, locale: str
) -> None:
    ctx = _context_for(world, corpus, locale)

    document = render_txt(ctx, output_dir=tmp_path)
    text = document.path.read_text(encoding="utf-8")

    assert FICTIONAL_NOTICE in text
    for block in ctx.recipe.blocks:
        assert block.text in text


@pytest.mark.parametrize("locale", _SAMPLE_LOCALES)
def test_xlsx_renderer_has_expected_sheets(
    world: SyntheticWorld, corpus: LocaleCorpus, tmp_path: Path, locale: str
) -> None:
    ctx = _context_for(world, corpus, locale)

    document = render_xlsx(ctx, output_dir=tmp_path)
    workbook = load_workbook(document.path)

    assert workbook.sheetnames == ["notice", "metrics", "blocks"]
    metrics_rows = list(workbook["metrics"].iter_rows(values_only=True))
    assert metrics_rows[0] == ("metric", "value")
    assert len(metrics_rows) > 1


@pytest.mark.parametrize("locale", _SAMPLE_LOCALES)
def test_docx_copyable_extracts_expected_text(
    world: SyntheticWorld, corpus: LocaleCorpus, tmp_path: Path, locale: str
) -> None:
    ctx = _context_for(world, corpus, locale)

    document = render_docx_copyable(ctx, output_dir=tmp_path)
    reader = DocxReader(str(document.path))
    full_text = "\n".join(p.text for p in reader.paragraphs)

    assert FICTIONAL_NOTICE in full_text
    for block in ctx.recipe.blocks:
        assert block.text in full_text


@pytest.mark.parametrize("locale", _SAMPLE_LOCALES)
def test_docx_non_copyable_has_no_text_layer(
    world: SyntheticWorld, corpus: LocaleCorpus, tmp_path: Path, locale: str
) -> None:
    ctx = _context_for(world, corpus, locale)

    document = render_docx_non_copyable(ctx, output_dir=tmp_path)
    reader = DocxReader(str(document.path))

    assert not any(p.text.strip() for p in reader.paragraphs)
    assert len(reader.inline_shapes) > 0


@pytest.mark.parametrize("locale", _SAMPLE_LOCALES)
def test_docx_mixed_has_both_text_and_an_image(
    world: SyntheticWorld, corpus: LocaleCorpus, tmp_path: Path, locale: str
) -> None:
    ctx = _context_for(world, corpus, locale)

    document = render_docx_mixed(ctx, output_dir=tmp_path)
    reader = DocxReader(str(document.path))

    assert any(p.text.strip() for p in reader.paragraphs)
    assert len(reader.inline_shapes) > 0


@pytest.mark.parametrize("locale", _SAMPLE_LOCALES)
def test_pdf_copyable_opens_and_extracts_non_empty_text(
    world: SyntheticWorld, corpus: LocaleCorpus, tmp_path: Path, locale: str
) -> None:
    ctx = _context_for(world, corpus, locale)

    document = render_pdf_copyable(ctx, output_dir=tmp_path)
    reader = PdfReader(document.path)
    text = " ".join(page.extract_text() for page in reader.pages)

    assert reader.pages
    assert FICTIONAL_NOTICE in text
    assert text.strip()


@pytest.mark.parametrize("locale", _SAMPLE_LOCALES)
def test_pdf_non_copyable_has_no_extractable_text(
    world: SyntheticWorld, corpus: LocaleCorpus, tmp_path: Path, locale: str
) -> None:
    ctx = _context_for(world, corpus, locale)

    document = render_pdf_non_copyable(ctx, output_dir=tmp_path)
    reader = PdfReader(document.path)
    text = " ".join(page.extract_text() for page in reader.pages).strip()

    assert reader.pages
    assert not text


@pytest.mark.parametrize("locale", _SAMPLE_LOCALES)
def test_pdf_mixed_has_a_copyable_page_and_an_image_page(
    world: SyntheticWorld, corpus: LocaleCorpus, tmp_path: Path, locale: str
) -> None:
    ctx = _context_for(world, corpus, locale)

    document = render_pdf_mixed(ctx, output_dir=tmp_path)
    reader = PdfReader(document.path)

    assert len(reader.pages) == 2
    assert reader.pages[0].extract_text().strip()


def test_every_one_of_the_15_locales_renders_a_full_pdf_with_glyphs_intact(
    world: SyntheticWorld, corpus: LocaleCorpus, tmp_path: Path
) -> None:
    """Full 15-locale sweep (unlike the 4-locale sample above): a copyable
    PDF is the most demanding format for font coverage, since a missing
    glyph in the embedded TTF makes reportlab raise instead of silently
    drawing a tofu box. Rendering every locale here is the regression test
    for the font_registry/text_shaping combination as a whole."""
    for pack in corpus.packs:
        ctx = _context_for(world, corpus, pack.locale)
        document = render_pdf_copyable(ctx, output_dir=tmp_path / pack.locale)
        reader = PdfReader(document.path)
        text = " ".join(page.extract_text() for page in reader.pages)
        assert text.strip(), f"{pack.locale}: no extractable text"
