from __future__ import annotations

import io
import random
from docx import Document
from pathlib import Path
from scripts.synthetic_data.core.models import FileFormat
from scripts.synthetic_data.core.random_content import (
    random_paragraphs,
    random_title,
)
from scripts.synthetic_data.core.text_image_renderer import render_text_snippet
from scripts.synthetic_data.generators.base import BaseGenerator


class DocxGenerator(BaseGenerator):
    """Mixes real editable paragraphs with one embedded picture-of-text paragraph in a single file."""

    file_format = FileFormat.DOCX
    extension = "docx"

    def _write(self, path: Path, rng: random.Random) -> None:
        document = Document()
        document.add_heading(random_title(rng), level=1)

        for paragraph in random_paragraphs(rng, count=rng.randint(2, 3)):
            document.add_paragraph(paragraph)

        document.add_heading("Scanned fragment", level=2)
        image_text = random_paragraphs(rng, count=1)[0]
        snippet_image = render_text_snippet(image_text)
        image_buffer = io.BytesIO()
        snippet_image.save(image_buffer, format="PNG")
        image_buffer.seek(0)
        document.add_picture(image_buffer)

        document.add_paragraph(random_paragraphs(rng, count=1)[0])

        document.save(str(path))
